import streamlit as st
import pandas as pd
import numpy as np
import tempfile
import os
import re
import base64
from PyPDF2 import PdfReader
from fpdf import FPDF
from io import BytesIO
# Use RapidFuzz, que é mais rápido para fuzzy matching
from rapidfuzz import process, fuzz

# Bibliotecas para gerar DOCX
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

###############################################################################
# FALLBACK PARA st.session_state (EVITA KeyError)
###############################################################################
_fallback_state = {
    "df_completo": None,
    "df_descontos": None,
    "df_descontos_gloss": None,
    "df_descontos_gloss_sel": None,
    "nome_cliente": None,
    "matricula": None,
    # Inserido para suportar valor B no cálculo de indébito:
    "valor_recebido": ""
}


def get_state_value(key):
    try:
        return st.session_state[key]
    except:
        return _fallback_state.get(key, None)


def set_state_value(key, value):
    try:
        st.session_state[key] = value
    except:
        _fallback_state[key] = value


###############################################################################
# CONFIGURAÇÃO INICIAL DO STREAMLIT
###############################################################################
st.set_page_config(page_title="Analista de Contracheques", layout="centered")

LOGO_PATH = "MP.png"  # Caminho para a logomarca
GLOSSARY_PATH = "Rubricas.txt"  # Nome do arquivo de Glossário (Rubricas.txt)


###############################################################################
# FUNÇÃO PARA SANITIZAR STRINGS (NOME, MATRICULA)
###############################################################################
def sanitizar_para_arquivo(texto: str) -> str:
    texto = texto.strip()
    texto = texto.replace(" ", "_")
    texto = re.sub(r"[^\w\-_\.]", "", texto, flags=re.UNICODE)
    return texto


###############################################################################
# FUNÇÃO PARA EXTRAIR NOME E MATRÍCULA (do PDF – não exibidos)
###############################################################################
def extrair_nome_e_matricula(pdf_path):
    nome = "N/D"
    matricula = "N/D"
    with open(pdf_path, 'rb') as f:
        reader = PdfReader(f)
        if len(reader.pages) > 0:
            text = reader.pages[0].extract_text() or ""
            lines = text.split('\n')
            for i, linha in enumerate(lines):
                if "NOME" in linha.upper():
                    if i + 1 < len(lines):
                        valor_nome = lines[i + 1].strip()
                        match_nome = re.match(r"([^\d]+)", valor_nome)
                        if match_nome:
                            nome = match_nome.group(1).strip()
                if "MATRÍCULA-SEQ-DIG" in linha.upper():
                    if i + 1 < len(lines):
                        valor_matr = lines[i + 1].strip()
                        matr_match = re.search(r"(\d{3}\.\d{3}-\d\s*[A-Z]*)", valor_matr)
                        if matr_match:
                            matricula = matr_match.group(1).strip()
    return nome or "N/D", matricula or "N/D"


###############################################################################
# FUNÇÃO PARA LIMPAR VALOR
###############################################################################
def limpar_valor(valor):
    if isinstance(valor, str):
        v = valor.replace(" ", "").replace(".", "").replace(",", ".")
        match_val = re.search(r"[\d\.]+", v)
        if match_val:
            return match_val.group(0)
    return valor


###############################################################################
# (1) ALTERAÇÃO DA FUNÇÃO DE INSERIR TOTAIS
#    AGORA COM 4 LINHAS:
#      A = Valor Total (R$)
#      B = Valor Recebido - Autor (a)
#      Indébito (A-B)
#      Indébito em dobro (R$)
###############################################################################
def inserir_totais_na_coluna(df, col_valor):
    """
    Antiga lógica: inseria "Valor Total (R$)" e "Em dobro (R$)".
    Agora insere 4 linhas:
      - A = Valor Total (R$)
      - B = Valor Recebido - Autor (a)
      - Indébito (A-B)
      - Indébito em dobro (R$)

    O valor B é recuperado da variável "valor_recebido" (session_state ou fallback).
    """
    if col_valor not in df.columns:
        return df

    def _to_float(x):
        try:
            return float(str(x).replace(',', '.').strip())
        except:
            return 0.0

    # Soma (A)
    soma = df[col_valor].apply(_to_float).sum()
    if soma == 0:
        return df

    df_novo = df.copy()

    # Recupera o valor B do estado
    valor_b_str = get_state_value("valor_recebido") or "0"
    try:
        valor_b_num = float(str(valor_b_str).replace(',', '.').strip())
    except:
        valor_b_num = 0.0

    # Calcula indebito e indebito em dobro
    indebito = soma - valor_b_num
    indebito_dobro = 2 * indebito

    def en_us_format(number: float) -> str:
        return f"{number:,.2f}"

    A_str = en_us_format(soma)
    B_str = valor_b_str.strip()
    indebito_str = en_us_format(indebito)
    indebito_dobro_str = en_us_format(indebito_dobro)

    # Linha A
    df_novo = pd.concat([
        df_novo,
        pd.DataFrame({col_valor: [A_str], "DESCRIÇÃO": ["A = Valor Total (R$)"]})
    ], ignore_index=True)
    # Linha B
    df_novo = pd.concat([
        df_novo,
        pd.DataFrame({col_valor: [B_str], "DESCRIÇÃO": ["B = Valor Recebido - Autor (a)"]})
    ], ignore_index=True)
    # Indébito (A-B)
    df_novo = pd.concat([
        df_novo,
        pd.DataFrame({col_valor: [indebito_str], "DESCRIÇÃO": ["Indébito (A-B)"]})
    ], ignore_index=True)
    # Indébito em dobro
    df_novo = pd.concat([
        df_novo,
        pd.DataFrame({col_valor: [indebito_dobro_str], "DESCRIÇÃO": ["Indébito em dobro (R$)"]})
    ], ignore_index=True)

    linhas_especiais = [
        "A = Valor Total (R$)",
        "B = Valor Recebido - Autor (a)",
        "Indébito (A-B)",
        "Indébito em dobro (R$)"
    ]
    mask_especial = df_novo["DESCRIÇÃO"].isin(linhas_especiais)
    if "DATA" in df_novo.columns:
        df_novo.loc[mask_especial, "DATA"] = ""
    if "COD" in df_novo.columns:
        df_novo.loc[mask_especial, "COD"] = ""

    return df_novo


###############################################################################
# FUNÇÕES GERAIS DE SUPORTE (Glossário e Imagens)
###############################################################################
def get_image_base64(file_path):
    if not os.path.exists(file_path):
        return ""
    with open(file_path, "rb") as img_file:
        return base64.b64encode(img_file.read()).decode()


# Utiliza o diretório atual para compor o caminho completo para Rubricas.txt
def carregar_glossario(path):
    try:
        full_path = os.path.join(os.getcwd(), path)
        with open(full_path, "r", encoding="utf-8") as f:
            return f.read().splitlines()
    except Exception as e:
        st.error(f"Erro ao carregar glossário: {e}")
        return []


###############################################################################
# EXTRAÇÃO DE TABELAS (CONTRACHEQUE) VIA CAMELOT
###############################################################################
def extrair_data_da_pagina(pdf_path, page_number):
    try:
        with open(pdf_path, 'rb') as f:
            reader = PdfReader(f)
            if page_number - 1 < len(reader.pages):
                text = reader.pages[page_number - 1].extract_text() or ""
                match = re.search(r"\d{2}/\d{4}", text)
                if match:
                    return match.group(0)
    except:
        pass
    return "N/D"


def _separar_linhas_multiplas(df: pd.DataFrame) -> pd.DataFrame:
    linhas_expandidas = []
    for _, row in df.iterrows():
        col_split = [str(row[col]).split('\n') for col in df.columns]
        max_splits = max(len(partes) for partes in col_split)
        for i in range(max_splits):
            nova_linha = {}
            for c, nome_coluna in enumerate(df.columns):
                partes = col_split[c]
                nova_linha[nome_coluna] = partes[i].strip() if i < len(partes) else ''
            linhas_expandidas.append(nova_linha)
    return pd.DataFrame(linhas_expandidas)


def encontrar_cabecalho(df):
    for idx, row in df.iterrows():
        if row.astype(str).str.contains(r"des[çc]rição", case=False, regex=True).any():
            return idx
    return None


def ler_tabelas(pdf_path):
    try:
        import camelot
        tables = camelot.read_pdf(
            pdf_path,
            pages="all",
            flavor="lattice",
            strip_text=''
        )
        if len(tables) == 0:
            tables = camelot.read_pdf(
                pdf_path,
                pages="all",
                flavor="stream",
                strip_text=''
            )
        return tables
    except Exception as e:
        st.error(f"Erro ao ler tabelas: {e}")
        return []


def ajustar_descontos_uma_pagina(df):
    discount_values = []
    for _, row in df.iterrows():
        d_val = str(row["DESCONTOS"]).strip()
        if d_val and d_val != "-":
            discount_values.append(d_val)
    last_ganhos_index = -1
    for i, row in enumerate(df.iterrows()):
        g_val = str(df.at[i, "GANHOS"]).strip()
        if g_val and g_val != "-" and re.search(r"\d", g_val):
            last_ganhos_index = i
        else:
            break
    start_index = last_ganhos_index + 1
    discount_index = 0
    for i in range(0, start_index):
        df.at[i, "DESCONTOS"] = ""
    for i in range(start_index, len(df)):
        if discount_index < len(discount_values):
            df.at[i, "DESCONTOS"] = discount_values[discount_index]
            discount_index += 1
        else:
            df.at[i, "DESCONTOS"] = ""
    return df


def ajustar_descontos_por_pagina(df):
    if "PAGINA" not in df.columns:
        return df
    paginas_processadas = []
    for page_number, group in df.groupby("PAGINA", group_keys=False):
        group = group.reset_index(drop=True)
        group_ajustado = ajustar_descontos_uma_pagina(group)
        group_ajustado["PAGINA"] = page_number
        paginas_processadas.append(group_ajustado)
    if not paginas_processadas:
        return df
    return pd.concat(paginas_processadas, ignore_index=True)


def processar_contracheque(pdf_path):
    colunas_desejadas = ["COD", "DESCRIÇÃO", "GANHOS", "DESCONTOS"]
    colunas_finais = colunas_desejadas + ["PAGINA", "DATA"]
    dados_finais = pd.DataFrame(columns=colunas_finais)
    tables = ler_tabelas(pdf_path)
    for table in tables:
        df = table.df
        idx_cab = encontrar_cabecalho(df)
        if idx_cab is None:
            continue
        df = df.iloc[idx_cab + 1:].reset_index(drop=True)
        if df.shape[1] >= 7:
            df = df.iloc[:, [0, 1, 5, 6]]
            df.columns = colunas_desejadas
        else:
            continue
        df = _separar_linhas_multiplas(df)
        for col in ["GANHOS", "DESCONTOS"]:
            df[col] = df[col].apply(limpar_valor)
        pagina_atual = table.page
        data_encontrada = extrair_data_da_pagina(pdf_path, pagina_atual)
        df["PAGINA"] = pagina_atual
        df["DATA"] = data_encontrada
        dados_finais = pd.concat([dados_finais, df], ignore_index=True)
    dados_finais.replace('', pd.NA, inplace=True)
    dados_finais.dropna(how='all', inplace=True)
    dados_finais.fillna('', inplace=True)
    dados_finais = ajustar_descontos_por_pagina(dados_finais)
    return dados_finais


###############################################################################
# FUNÇÕES PARA GERAÇÃO DE PDF E DOCX (mantidas inalteradas, exceto pela
# chamada a inserir_totais_na_coluna que agora gera as 4 linhas solicitadas)
###############################################################################
def formatar_valor_brl(us_string: str) -> str:
    try:
        f = float(us_string.replace(",", "").replace(".", "")) / 100
        return f"{f:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except:
        return us_string


class PDFRelatorio(FPDF):
    def __init__(self, titulo, colunas, dados, linhas_especiais=False):
        super().__init__(orientation='L', unit='mm', format='A4')
        self.titulo = titulo
        self.colunas = colunas
        self.dados = dados
        self.linhas_especiais = linhas_especiais
        self.set_auto_page_break(auto=False, margin=15)
        self.set_left_margin(10)
        self.set_right_margin(10)
        self.set_top_margin(10)

    def header(self):
        self.set_font('Arial', 'B', 14)
        self.cell(0, 8, self.titulo, border=False, ln=True, align='C')
        self.ln(3)
        self.set_font("Arial", "B", 10)
        self.set_fill_color(200, 220, 255)
        for col in self.colunas:
            self.cell(col["largura"], 8, col["nome"], border=1, align='C', fill=True)
        self.ln()

    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.cell(0, 10, f'Página {self.page_no()}', border=False, ln=False, align='C')

    def montar_tabela(self):
        self.set_font("Arial", "", 9)
        row_height = 7
        for _, row in self.dados.iterrows():
            if self.get_y() + row_height + 15 > self.h:
                self.add_page()

            descricao = str(row.get("DESCRIÇÃO", ""))
            # Ajustamos para as novas linhas especiais
            linhas_quentes = [
                "A = Valor Total (R$)",
                "B = Valor Recebido - Autor (a)",
                "Indébito (A-B)",
                "Indébito em dobro (R$)"
            ]
            is_especial = (descricao in linhas_quentes)

            if is_especial and self.linhas_especiais:
                self.set_font("Arial", "B", 11)
                self.set_text_color(255, 0, 0)
            else:
                self.set_font("Arial", "", 9)
                self.set_text_color(0, 0, 0)

            for col in self.colunas:
                col_name = col["nome"]
                valor = str(row.get(col_name, ""))
                if col_name in ["GANHOS", "DESCONTOS"] and valor.strip():
                    valor = formatar_valor_brl(valor)
                self.cell(col["largura"], row_height, valor, border=1, align=col["alinhamento"])
            self.ln(row_height)

            if is_especial and self.linhas_especiais:
                self.set_font("Arial", "", 9)
                self.set_text_color(0, 0, 0)

    def gerar_pdf(self, nome_arquivo):
        self.add_page()
        self.montar_tabela()
        self.output(nome_arquivo)


def salvar_em_pdf(dados: pd.DataFrame, titulo_pdf: str, colunas_def: list,
                  inserir_totais=False, col_valor_soma="DESCONTOS",
                  linhas_especiais=False) -> bytes:
    for col_def in colunas_def:
        if col_def["nome"] not in dados.columns:
            dados[col_def["nome"]] = ""
    df_final = dados.copy()
    if inserir_totais:
        df_final = inserir_totais_na_coluna(df_final, col_valor_soma)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
        tmp_path = tmp_pdf.name
    pdf = PDFRelatorio(titulo_pdf, colunas_def, df_final, linhas_especiais=linhas_especiais)
    pdf.gerar_pdf(tmp_path)
    with open(tmp_path, "rb") as f:
        pdf_bytes = f.read()
    os.remove(tmp_path)
    return pdf_bytes


def to_en_us_string(val):
    try:
        f = float(str(val).replace(",", "."))
        return "{:,.2f}".format(f)
    except:
        return str(val)


def df_to_docx_bytes(dados: pd.DataFrame, titulo: str,
                     inserir_totais=False, col_valor_soma="DESCONTOS") -> bytes:
    df_final = dados.copy()
    if inserir_totais:
        df_final = inserir_totais_na_coluna(df_final, col_valor_soma)
    document = Document()
    for section in document.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height
    titulo_heading = document.add_heading(titulo, level=1)
    titulo_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if df_final.empty:
        p = document.add_paragraph("DataFrame vazio - nenhum dado para exibir.")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        buf = BytesIO()
        document.save(buf)
        return buf.getvalue()

    colunas = df_final.columns.tolist()
    table = document.add_table(rows=1, cols=len(colunas))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(colunas):
        hdr_cells[i].text = str(col_name)
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    width_map = {}
    if "COD" in colunas:
        width_map["COD"] = 20
    if "DESCRIÇÃO" in colunas:
        width_map["DESCRIÇÃO"] = 130
    if "GANHOS" in colunas:
        width_map["GANHOS"] = 40
    if "DESCONTOS" in colunas:
        width_map["DESCONTOS"] = 40
    if "PAGINA" in colunas:
        width_map["PAGINA"] = 20
    if "DATA" in colunas:
        width_map["DATA"] = 30

    # Linhas da Tabela
    linhas_quentes = [
        "A = Valor Total (R$)",
        "B = Valor Recebido - Autor (a)",
        "Indébito (A-B)",
        "Indébito em dobro (R$)"
    ]
    for _, row in df_final.iterrows():
        descricao = str(row.get("DESCRIÇÃO", ""))
        is_especial = (descricao in linhas_quentes)

        row_cells = table.add_row().cells
        for i, col_name in enumerate(colunas):
            valor = str(row[col_name])
            if col_name in ["GANHOS", "DESCONTOS"] and valor.strip():
                valor = to_en_us_string(valor)
            paragraph = row_cells[i].paragraphs[0]
            run = paragraph.add_run(valor)
            if col_name.upper() == "DESCRIÇÃO":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run.font.size = Pt(9)
            if is_especial:
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(255, 0, 0)

    for i, col_name in enumerate(colunas):
        mm = width_map.get(col_name, 25)
        table.columns[i].width = Inches(mm / 25.4)

    buf = BytesIO()
    document.save(buf)
    return buf.getvalue()


def formatar_valor_brl(valor):
    try:
        f = float(str(valor).replace(",", "").replace(".", "")) / 100
        return f"{f:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except:
        return str(valor)


def ajustar_valores_docx(file_input_bytes: bytes) -> bytes:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_in:
        tmp_in.write(file_input_bytes)
        tmp_in.flush()
        input_path = tmp_in.name
    output_path = input_path.replace(".docx", "_corrigido.docx")
    doc = Document(input_path)
    pattern = re.compile(r'([\d,]+\.\d{2})')
    for para in doc.paragraphs:
        found = pattern.findall(para.text)
        if not found:
            continue
        for val_us in found:
            val_br = formatar_valor_brl(val_us)
            para.text = para.text.replace(val_us, val_br)
    doc.save(output_path)
    with open(output_path, "rb") as f:
        final_bytes = f.read()
    os.remove(input_path)
    os.remove(output_path)
    return final_bytes


###############################################################################
# Função para cruzar o Extrato de Descontos com a Lista de Rubricas
###############################################################################
def cruzar_descontos_com_rubricas(df_descontos, glossary, threshold=85):
    if df_descontos.empty or not glossary:
        return pd.DataFrame()
    unique_desc = df_descontos["DESCRIÇÃO"].unique()
    mapping = {}
    for desc in unique_desc:
        result = process.extractOne(desc, glossary, scorer=fuzz.ratio)
        mapping[desc] = (result is not None and result[1] >= threshold)
    mask = df_descontos["DESCRIÇÃO"].map(mapping)
    return df_descontos[mask]


###############################################################################
# APLICAÇÃO STREAMLIT (MAIN)
###############################################################################
def main():
    # Exibir logomarca
    logo_base64 = get_image_base64(LOGO_PATH)
    if logo_base64:
        st.markdown(f"""
            <div style="text-align: center; margin-bottom: 20px;">
                <img src="data:image/png;base64,{logo_base64}" alt="Logomarca" style="width: 300px;">
            </div>
            """, unsafe_allow_html=True)

    st.title("Analista de Contracheques")

    # Carregar glossário (lista de Rubricas)
    glossary_terms = carregar_glossario(GLOSSARY_PATH)

    # Upload do PDF
    uploaded_pdf = st.file_uploader(
        "Clique no botão para enviar o arquivo PDF (Contracheque) - SEAD (com colunas GANHOS e DESCONTOS)",
        type="pdf"
    )
    if uploaded_pdf is not None:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(uploaded_pdf.read())
            caminho_temp = tmp.name

        nome_cli, matr = extrair_nome_e_matricula(caminho_temp)
        set_state_value("nome_cliente", nome_cli)
        set_state_value("matricula", matr)

        df = processar_contracheque(caminho_temp)
        os.unlink(caminho_temp)
        if not df.empty:
            set_state_value("df_completo", df)
        else:
            st.warning("Não foi possível extrair as informações do PDF ou o arquivo está vazio.")

    df_completo = get_state_value("df_completo")
    nome_cli_sanit = sanitizar_para_arquivo(get_state_value("nome_cliente") or "ND")
    matr_sanit = sanitizar_para_arquivo(get_state_value("matricula") or "ND")

    if df_completo is not None and not df_completo.empty:
        st.markdown("### DataFrame do Contracheque Completo")
        st.dataframe(df_completo, use_container_width=True)

        # Item 1: PDF Completo
        titulo_completo = f"Relatório de Contracheque (Completo) - {get_state_value('nome_cliente')} / {get_state_value('matricula')}"
        colunas_pdf_completo = [
            {"nome": "COD", "largura": 20, "alinhamento": "C"},
            {"nome": "DESCRIÇÃO", "largura": 140, "alinhamento": "L"},
            {"nome": "GANHOS", "largura": 30, "alinhamento": "R"},
            {"nome": "DESCONTOS", "largura": 30, "alinhamento": "R"},
            {"nome": "PAGINA", "largura": 20, "alinhamento": "C"},
            {"nome": "DATA", "largura": 30, "alinhamento": "C"},
        ]
        pdf_data_completo = salvar_em_pdf(
            dados=df_completo.copy(),
            titulo_pdf=titulo_completo,
            colunas_def=colunas_pdf_completo,
            inserir_totais=False,
            col_valor_soma="DESCONTOS",
            linhas_especiais=False
        )
        pdf_filename_completo = f"contracheque_completo_{nome_cli_sanit}_{matr_sanit}.pdf"
        st.download_button(
            label="Baixar PDF (Completo)",
            data=pdf_data_completo,
            file_name=pdf_filename_completo,
            mime="application/pdf"
        )

        st.markdown("## Análise de Descontos")

        # Form para filtrar Descontos
        with st.form("form_filtrar_descontos"):
            st.markdown("### 1) Filtrar Operações de Descontos")
            submit_desc = st.form_submit_button("Filtrar Descontos")
        if submit_desc:
            df_desc = df_completo.drop(columns=["GANHOS"], errors='ignore')
            df_desc = df_desc[df_desc["DESCONTOS"].str.strip() != ""]
            df_desc.reset_index(drop=True, inplace=True)
            set_state_value("df_descontos", df_desc)

        df_descontos = get_state_value("df_descontos")
        if df_descontos is not None and not df_descontos.empty:
            st.markdown("### 2) Extrato de Descontos")
            st.dataframe(df_descontos, use_container_width=True)

            # Botão de Baixar PDF (Descontos)
            titulo_desc = f"Contracheque - Descontos - {get_state_value('nome_cliente')} / {get_state_value('matricula')}"
            colunas_pdf_desc = [
                {"nome": "COD", "largura": 20, "alinhamento": "C"},
                {"nome": "DESCRIÇÃO", "largura": 160, "alinhamento": "L"},
                {"nome": "DESCONTOS", "largura": 30, "alinhamento": "R"},
                {"nome": "PAGINA", "largura": 20, "alinhamento": "C"},
                {"nome": "DATA", "largura": 30, "alinhamento": "C"},
            ]
            pdf_data_desc = salvar_em_pdf(
                dados=df_descontos.copy(),
                titulo_pdf=titulo_desc,
                colunas_def=colunas_pdf_desc,
                inserir_totais=False,
                col_valor_soma="DESCONTOS",
                linhas_especiais=False
            )
            pdf_filename_desc = f"contracheque_descontos_{nome_cli_sanit}_{matr_sanit}.pdf"
            st.download_button(
                label="Baixar PDF (Descontos)",
                data=pdf_data_desc,
                file_name=pdf_filename_desc,
                mime="application/pdf"
            )

            # (2.1) Lista das Rubricas
            st.markdown("### 2.1) Lista das Rubricas")
            df_rubricas = pd.DataFrame({"Rubricas": glossary_terms})
            st.dataframe(df_rubricas, use_container_width=True)

            # (3) Cruzamento entre Extrato de Descontos e Rubricas
            with st.form("form_filtro_gloss"):
                st.markdown("### 3) Filtrar Descontos no Glossário (Precisão Ajustável)")
                thresh = st.slider("Nível de Similaridade (0.1 a 1.0)", 0.1, 1.0, 0.85, 0.1)
                submit_gloss = st.form_submit_button("Filtrar com Rubricas")
            if submit_gloss:
                with st.spinner("Cruzando Extrato de Descontos com a Lista das Rubricas..."):
                    threshold_value = int(thresh * 100)
                    df_desc_gloss = cruzar_descontos_com_rubricas(df_descontos, glossary_terms, threshold_value)
                set_state_value("df_descontos_gloss", df_desc_gloss)
                set_state_value("df_descontos_gloss_sel", None)

        df_descontos_gloss = get_state_value("df_descontos_gloss")
        if df_descontos_gloss is not None and not df_descontos_gloss.empty:
            st.markdown("#### Descontos x Glossário")
            st.dataframe(df_descontos_gloss, use_container_width=True)
            titulo_gloss = f"Descontos x Glossário - {get_state_value('nome_cliente')} / {get_state_value('matricula')}"
            colunas_pdf_gloss = [
                {"nome": "COD", "largura": 20, "alinhamento": "C"},
                {"nome": "DESCRIÇÃO", "largura": 160, "alinhamento": "L"},
                {"nome": "DESCONTOS", "largura": 30, "alinhamento": "R"},
                {"nome": "PAGINA", "largura": 20, "alinhamento": "C"},
                {"nome": "DATA", "largura": 30, "alinhamento": "C"},
            ]
            pdf_data_gloss = salvar_em_pdf(
                dados=df_descontos_gloss.copy(),
                titulo_pdf=titulo_gloss,
                colunas_def=colunas_pdf_gloss,
                inserir_totais=False,
                col_valor_soma="DESCONTOS",
                linhas_especiais=False
            )
            pdf_filename_gloss = f"contracheque_descontos_glossario_{nome_cli_sanit}_{matr_sanit}.pdf"
            st.download_button(
                label="Baixar PDF (Descontos x Glossário)",
                data=pdf_data_gloss,
                file_name=pdf_filename_gloss,
                mime="application/pdf"
            )

            # (4) Lista única de Descontos
            df_gloss_origem = df_descontos_gloss
            df_sel = get_state_value("df_descontos_gloss_sel")
            if df_sel is None or df_sel.empty:
                df_sel = df_gloss_origem
            with st.form("form_inclusao_descontos"):
                st.markdown("### 4) Lista Única de Descontos")
                valores_unicos = sorted(df_sel["DESCRIÇÃO"].unique())
                st.write("Marque os itens que deseja incluir:")
                selected_descr = []
                for i, val in enumerate(valores_unicos):
                    qtd = df_sel[df_sel["DESCRIÇÃO"] == val].shape[0]
                    label_str = f"{i + 1} - {val} ({qtd}x)"
                    if st.checkbox(label_str, key=f"chk_{i}"):
                        selected_descr.append(val)
                incluir_btn = st.form_submit_button("Confirmar Inclusão (Descontos)")

            if incluir_btn:
                if selected_descr:
                    df_incluido = df_sel[df_sel["DESCRIÇÃO"].isin(selected_descr)].copy()
                    set_state_value("df_descontos_gloss_sel", df_incluido)
                    st.success("Descontos selecionados com sucesso!")
                    st.markdown("#### Lista Restante após Inclusões")
                    st.dataframe(df_incluido, use_container_width=True)
                else:
                    st.warning("Nenhuma descrição selecionada.")

            # (5) APRESENTAR RÚBRICAS PARA DÉBITOS (DESCONTOS FINAIS)
            df_final_sel = get_state_value("df_descontos_gloss_sel")
            if df_final_sel is not None and not df_final_sel.empty:
                ######################################################################
                # INSERINDO A ETAPA "Apresentar Rúbricas para Débitos (Descontos Finais)"
                # COM OS CAMPOS:
                #    A = Valor Total (R$)
                #    B = Valor Recebido - Autor (a)
                #    Indébito (A-B)
                #    Indébito em dobro (R$)
                ######################################################################

                st.markdown("### 5) Apresentar Rúbricas para Débitos (Descontos Finais)")

                # Cópia e ordenação cronológica
                df_final = df_final_sel.copy()
                df_final["PAGINA"] = pd.to_numeric(df_final["PAGINA"], errors='coerce').fillna(0)
                df_final = df_final.sort_values(by=["DATA", "PAGINA"]).reset_index(drop=True)
                df_final = df_final[["COD", "DESCRIÇÃO", "DESCONTOS", "DATA"]]

                # Cálculo de A (soma dos descontos)
                def _to_float(x):
                    try:
                        return float(str(x).replace(',', '.').strip())
                    except:
                        return 0.0

                A_val = df_final["DESCONTOS"].apply(_to_float).sum()
                A_str = f"{A_val:,.2f}"

                st.write(f"A = Valor Total (R$): {A_str}")

                col1, col2 = st.columns(2)
                with col1:
                    valor_b_receb = st.text_input("B = Valor Recebido - Autor (a)", "0")
                try:
                    vrnum = float(valor_b_receb.replace(',', '.').strip())
                except:
                    vrnum = 0.0

                indebito = A_val - vrnum
                indebito_dobro = 2 * indebito
                indebito_str = f"{indebito:,.2f}"
                indebito_dobro_str = f"{indebito_dobro:,.2f}"

                with col2:
                    st.write(f"Indébito (A-B): {indebito_str}")
                    st.write(f"Indébito em dobro (R$): {indebito_dobro_str}")

                # Armazena "valor_recebido" no estado
                set_state_value("valor_recebido", valor_b_receb)

                with st.form("form_descontos_finais"):
                    submit_final = st.form_submit_button("Gerar Relatório Final de Descontos")

                if submit_final:
                    # Monta Título final
                    nome = get_state_value("nome_cliente") or "ND"
                    matr_ = get_state_value("matricula") or "ND"
                    titulo_final = f"Descontos Finais (Cronológico) - {nome} / {matr_}"

                    # Insere 4 linhas especiais (A, B, Indébito, Indébito em dobro)
                    df_com_totais = inserir_totais_na_coluna(df_final.copy(), "DESCONTOS")

                    # Gera PDF final
                    colunas_pdf_finais = [
                        {"nome": "COD", "largura": 20, "alinhamento": "C"},
                        {"nome": "DESCRIÇÃO", "largura": 180, "alinhamento": "L"},
                        {"nome": "DESCONTOS", "largura": 30, "alinhamento": "R"},
                        {"nome": "DATA", "largura": 30, "alinhamento": "C"},
                    ]
                    pdf_data_finais = salvar_em_pdf(
                        dados=df_com_totais,
                        titulo_pdf=titulo_final,
                        colunas_def=colunas_pdf_finais,
                        inserir_totais=False,     # Já inserimos manualmente
                        col_valor_soma="DESCONTOS",
                        linhas_especiais=True     # Destaca as 4 linhas
                    )
                    pdf_filename_finais = f"contracheque_descontos_finais_{nome_cli_sanit}_{matr_sanit}.pdf"
                    st.download_button(
                        label="Baixar PDF (Descontos Finais - Cronológico)",
                        data=pdf_data_finais,
                        file_name=pdf_filename_finais,
                        mime="application/pdf"
                    )

                    # Gera DOCX final
                    docx_bytes = df_to_docx_bytes(
                        dados=df_final.copy(),    # note: gera sem tot, mas passamos "inserir_totais=True"
                        titulo=titulo_final,
                        inserir_totais=True,      # inc. A, B, Indébito, Indébito em dobro
                        col_valor_soma="DESCONTOS"
                    )
                    docx_bytes_corrigido = ajustar_valores_docx(docx_bytes)
                    docx_filename_finais = pdf_filename_finais.replace(".pdf", ".docx")
                    st.download_button(
                        label="Baixar DOCX (Descontos Finais - Cronológico)",
                        data=docx_bytes_corrigido,
                        file_name=docx_filename_finais,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )


if __name__ == "__main__":
    main()




